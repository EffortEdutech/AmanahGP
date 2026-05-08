from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path("docs/pitch/2026-05-usm")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(9)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], "1F6F5B")
        if widths:
            hdr[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def build_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 111, 91)
    styles["Heading 2"].font.name = "Aptos"
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.color.rgb = RGBColor(35, 45, 57)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Digitising Amanah")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(23, 83, 69)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "A behavioural and digital governance pilot for Malaysian Islamic social finance organizations"
    ).italic = True

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    meta_rows = [
        ("Applicant", "Amanah Governance Platform / Effort Edutech"),
        ("Applicant structure", "AGP direct applicant; Prof. Anuar as advising member; JAKIM/MAIN as proposed development partner"),
        ("Funding request", "RM550,000 for a 12-month pilot"),
        ("Pilot scale", "10-20 mosques, Islamic NGOs, waqf initiatives, or community charities"),
    ]
    for row, (k, v) in zip(meta.rows, meta_rows):
        set_cell_text(row.cells[0], k, bold=True)
        set_cell_shading(row.cells[0], "E7F4EF")
        set_cell_text(row.cells[1], v)
    doc.add_paragraph()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Amanah Governance Platform (AGP) provides a digital trust layer for Islamic charities, mosques, "
        "waqf initiatives and community organizations. It combines public charity transparency, organization-side "
        "governance operations, evaluator workflows, and an auditable Amanah Index. AGP is designed as public-interest "
        "infrastructure rather than a commercial product, and charity organizations will not pay AGP for the services "
        "they receive. This pilot proposes AGP as direct applicant, Prof. Anuar as advising member, and JAKIM/MAIN as "
        "grant, development or implementation partners. The advisory layer grounds the project in organizational behaviour, "
        "HRD, organizational learning, ethics, and community knowledge transfer."
    )
    doc.add_paragraph(
        "AGP is already live with seeded/mock public data. AmanahHub's charity directory currently demonstrates 46 visible "
        "organisations and 37 published trust profiles across mosque/surau, waqf institution, zakat body, foundation/yayasan, "
        "cooperative, welfare association and other categories. This gives funders and institutional partners a working MVP to inspect, "
        "not only a concept proposal."
    )

    doc.add_heading("The Problem", level=1)
    for item in [
        "Grassroots organizations often manage evidence, reports and compliance records through scattered documents and manual processes.",
        "Donors cannot easily compare governance quality or see whether an organization is improving over time.",
        "Reviewers and scholars need structured workflows for clarifications, approvals and publication control.",
        "Digital adoption is a behavioural challenge: administrators need confidence, training, routines and leadership reinforcement.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("The Solution", level=1)
    add_table(
        doc,
        ["Layer", "Function", "Pilot value"],
        [
            ("AmanahHub", "Public donor trust profiles and non-custodial giving", "Makes verified governance visible to donors"),
            ("amanahOS", "Organization workspace for governance, accounting, reports and certification", "Builds administrator reporting discipline"),
            ("AGP Console", "Reviewer, scholar and platform authority workflow", "Creates auditable review and clarification loops"),
            ("Amanah Index", "Weighted trust scoring engine", "Turns governance behaviour into transparent signals"),
        ],
    )

    doc.add_heading("Live MVP Evidence", level=1)
    for item in [
        "AmanahHub public charity directory with 46 visible organisations.",
        "37 published trust profiles with Amanah scores and certification tiers.",
        "Organisation filters for NGO/welfare association, mosque/surau, waqf institution, zakat body, foundation/yayasan, cooperative and other.",
        "State filters across Malaysia.",
        "Seeded/mock profiles covering waqf dialysis support, waqf clinic support, hospital equipment funds, zakat bodies, education foundations and mosque/surau programmes.",
        "Live amanahOS login surface for charity/operator workspace.",
        "Live AGP Console login surface for platform, reviewer and administrative workflows.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Advisory and Institutional Partner Roles", level=1)
    doc.add_paragraph(
        "Prof. Anuar will be invited as an advising member rather than making USM the formal applicant anchor. His role can "
        "strengthen the behavioural research model, training design, adoption measurement and evaluation of whether AGP changes "
        "real institutional practice. This aligns with his publicly listed expertise in Organizational Behaviour, HRM, "
        "Organizational Development, HRD, Organizational Learning, Business Ethics and community knowledge transfer."
    )
    doc.add_paragraph(
        "JAKIM and state MAINs will be approached as grant, development or implementation partners because they are the "
        "natural authority-side beneficiaries. AGP can register organizations under their authority, including mosques, "
        "suraus, Islamic charities, waqf initiatives and other approved entities. Organization members can then update "
        "profiles, reports, evidence, bookkeeping records and certification progress over time."
    )

    doc.add_heading("Objectives", level=1)
    for item in [
        "Deploy AGP with a JAKIM/MAIN pilot cohort.",
        "Register the agreed authority-side organizations, including mosques, suraus and Islamic charities.",
        "Improve evidence completeness, report timeliness and remediation closure.",
        "Measure whether visible trust scores improve donor confidence.",
        "Develop a practical capacity-building toolkit for mosque and NGO administrators.",
        "Produce a research-backed scale blueprint for Malaysian Islamic social finance governance.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Methodology", level=1)
    add_table(
        doc,
        ["Phase", "Activities", "Outputs"],
        [
            ("Setup", "Confirm advisor, approach JAKIM/MAIN, define authority scope", "Pilot protocol, advisor terms, partner proposal"),
            ("Baseline + Training", "Surveys, interviews, onboarding workshops", "Baseline report, trained administrators"),
            ("Registry + Deployment", "Register organizations, activate accounts, evidence upload, trust events", "Registered cohort, usage data, review records"),
            ("Improvement Cycle", "Clarifications, remediation, follow-up training", "Remediation data, revised toolkit"),
            ("Evaluation", "Endline analysis, grant reporting, publication planning", "Final report, paper outline, scale roadmap"),
        ],
    )

    doc.add_heading("KPIs", level=1)
    add_table(
        doc,
        ["KPI", "Target"],
        [
            ("Pilot organizations onboarded", "10-20"),
            ("Authority-side registry coverage", "100% of agreed JAKIM/MAIN pilot cohort"),
            ("Core evidence pack completion", "80%"),
            ("Timely report submission", "50% improvement"),
            ("Amanah score movement", "+10 to +15 points where baseline gaps exist"),
            ("Review findings closed", "70% within agreed cycle"),
            ("Donor trust perception", "20% improvement"),
            ("Training satisfaction", "85% positive"),
            ("Research/practitioner outputs", "1 report, 1 toolkit, 1 manuscript or conference paper"),
        ],
    )

    doc.add_heading("Budget", level=1)
    add_table(
        doc,
        ["Item", "Amount (RM)", "Notes"],
        [
            ("Platform development and configuration", "180,000", "Pilot features, data exports, trust profile polish"),
            ("Research design and evaluation", "90,000", "Instruments, interviews, analysis, reporting"),
            ("Training and capacity building", "80,000", "Workshops, materials, facilitation"),
            ("Pilot operations and support", "90,000", "Helpdesk, field coordination, onboarding"),
            ("Cloud, security and compliance", "55,000", "Hosting, storage, monitoring, backup"),
            ("Communications and dissemination", "25,000", "Reports, toolkit, stakeholder roundtable"),
            ("Contingency", "30,000", "Delivery and field risk buffer"),
            ("Total", "550,000", "12-month pilot"),
        ],
    )

    doc.add_heading("Sustainability", level=1)
    doc.add_paragraph(
        "AGP's long-term sustainability model is deliberately structured so that charity organizations do not pay AGP. "
        "The platform is intended as public-interest infrastructure. Initial development, expansion and operational readiness "
        "will be supported through grants. Governance onboarding programmes can be sponsored by corporate ESG/CSR partners. "
        "Long-term continuity will come from alignment with zakat and waqf ecosystems, where institutions benefit from stronger "
        "governance, audit-readiness, reporting consistency and sector-wide trust data."
    )
    doc.add_paragraph(
        "AGP may charge zakat bodies, waqf institutions, foundations, councils, CSR funders or other institutional sponsors "
        "for service packages delivered to charity organizations. These sponsored services include:"
    )
    for item in [
        "Organisation SaaS subscription.",
        "Managed bookkeeping service.",
        "Audit-readiness programme.",
        "Certification support.",
        "Institutional dashboard subscription.",
        "CSR/funder reporting tools.",
        "Training and governance workshops.",
        "Reviewer and Shariah marketplace.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("JAKIM/MAIN Development Partner Value", level=1)
    for item in [
        "JAKIM/MAIN obtain a living registry of organizations under their authority.",
        "AGP handles initial platform registration for mosques, suraus and other approved entities.",
        "Organization members update profiles, reports, projects, evidence, bookkeeping records and certification progress over time.",
        "Institutional dashboards provide visibility into reporting gaps, audit-readiness, risk signals and improvement progress.",
        "Charity organizations receive onboarding and governance support without paying AGP.",
        "Donors and the public gain confidence through verified AmanahHub profiles where publication is approved.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Diversified Grant and Partnership Pipeline", level=1)
    doc.add_paragraph(
        "AGP should not rely on a single funder or single institutional route. The JAKIM/MAIN pathway is strategically important, "
        "but all viable routes remain active."
    )
    for item in [
        "Malaysia Digital (MD) Status to unlock MDEC MDCG and MDAG pathways.",
        "Hasanah Micro Grant and Hasanah Special Grant for pilot and multi-year community impact.",
        "Cradle CIP Spark / Sprint for MVP validation and commercialization readiness.",
        "SC FIKRA ACE for Islamic fintech mentorship, visibility and industry connections.",
        "JAWHAR, JAKIM and state MAINs for institutional anchoring and authority-side deployment.",
        "IERIF / INCEIF for Islamic economics, governance, AI/data and social finance research.",
        "IsDB Engage, Transform Fund, Tadamon and later OIC-region scale opportunities.",
        "Corporate CSR/ESG sponsorships for governance onboarding and audit-readiness cohorts.",
        "State-level programmes linked to Islamic affairs, waqf, zakat, digital economy and social impact.",
        "University-linked co-applications where the funder requires or rewards academic anchoring.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Closing Case", level=1)
    doc.add_paragraph(
        "AGP gives Malaysia a practical path to digitise amanah. With AGP as direct applicant, JAKIM/MAIN as development partners, "
        "Prof. Anuar as advising member, and a diversified grant pipeline kept open, the project can move quickly while still proving "
        "that the platform changes governance behaviour in the institutions that hold community trust."
    )

    out = OUT / "AGP_Grant_Application_USM_Partner_Draft.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print(build_docx())
